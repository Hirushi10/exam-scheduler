import streamlit as st
import pandas as pd
from datetime import datetime, time
import docx
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os
import pymysql

st.set_page_config(page_title="Exams Scheduler", layout="wide")

# Custom CSS for the desktop app design
st.markdown("""
    <style>
    .main-title { font-size:26px; font-weight:bold; color: #1E3A8A; text-align: left; margin-bottom: 15px; border-bottom: 2px solid #1E3A8A; padding-bottom: 5px; }
    .hall-box { border: 1px solid #CBD5E1; padding: 15px; background-color: #FFFFFF; border-radius: 5px; margin-bottom: 15px; min-height: 240px; box-shadow: 1px 1px 5px rgba(0,0,0,0.05); }
    .hall-title { font-size: 16px; font-weight: bold; color: #1E40AF; border-bottom: 1px solid #E2E8F0; padding-bottom: 3px; margin-bottom: 8px; }
    .hall-details { font-size: 13px; color: #334155; line-height: 1.6; }
    .sidebar-section { background-color: #F8FAFC; padding: 15px; border-radius: 5px; border: 1px solid #E2E8F0; margin-bottom: 15px; }
    .summary-title { font-size: 15px; font-weight: bold; color: #0F172A; margin-bottom: 10px; border-bottom: 1px solid #CBD5E1; }
    .badge-capacity { font-size: 11px; font-weight: bold; background-color: #E2E8F0; color: #475569; padding: 3px 6px; border-radius: 4px; float: right; }
    .badge-danger { font-size: 11px; font-weight: bold; background-color: #EF4444; color: #FFFFFF; padding: 3px 6px; border-radius: 4px; float: right; }
    .badge-success { font-size: 11px; font-weight: bold; background-color: #10B981; color: #FFFFFF; padding: 3px 6px; border-radius: 4px; float: right; }
    
    /* Force Hand Pointer Symbol on Dropdowns, Inputs, and Selections */
    div[data-baseweb="select"] { cursor: pointer !important; }
    div[data-baseweb="select"] * { cursor: pointer !important; }
    div[data-testid="stDateInput"] input { cursor: pointer !important; }
    div[data-testid="stTimeInput"] input { cursor: pointer !important; }
    button { cursor: pointer !important; }
    </style>
""", unsafe_allow_html=True)

st.markdown("<div class='main-title'>💻 Exams Scheduler - Faculty Of Management & Finance</div>", unsafe_allow_html=True)

# ----------------- DATABASE CONNECTION -----------------
def get_db_connection():
    return pymysql.connect(
        host=st.secrets["mysql"]["host"],
        port=int(st.secrets["mysql"]["port"]),
        database=st.secrets["mysql"]["database"],
        user=st.secrets["mysql"]["user"],
        password=st.secrets["mysql"]["password"],
        autocommit=True,
        connect_timeout=10
    )

def init_db():
    conn = get_db_connection()
    cursor = conn.cursor()
    cursor.execute("""
        CREATE TABLE IF NOT EXISTS exam_assignments (
            id INT AUTO_INCREMENT PRIMARY KEY,
            semester_key VARCHAR(50),
            date VARCHAR(20),
            time VARCHAR(50),
            subject VARCHAR(255),
            hall VARCHAR(255),
            supervisor VARCHAR(255),
            invigilators TEXT
        )
    """)
    
    cursor.execute("SHOW COLUMNS FROM exam_assignments LIKE 'student_count'")
    if not cursor.fetchone():
        cursor.execute("ALTER TABLE exam_assignments ADD COLUMN student_count INT DEFAULT 0")
        
    cursor.close()
    conn.close()

try:
    init_db()
except Exception as e:
    st.error(f"Database Initialization Error: {e}")
    st.stop()

def load_all_schedules():
    conn = get_db_connection()
    query = "SELECT id, date AS Date, time AS Time, subject AS Subject, hall AS Hall, supervisor AS Supervisor, invigilators AS Invigilators, semester_key, student_count FROM exam_assignments"
    df = pd.read_sql(query, conn)
    conn.close()
    return df

schedule_df = load_all_schedules()

# Word Document Generation Function
def generate_word_report(df, title_text):
    doc = docx.Document()
    title = doc.add_paragraph()
    title_run = title.add_run(f"UNIVERSITY OF RUHUNA\nFaculty of Management and Finance\n{title_text}")
    title_run.bold = True
    title_run.font.size = Pt(14)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("\n")
    
    cols_to_drop = [c for c in ['id', 'semester_key'] if c in df.columns]
    report_df = df.drop(columns=cols_to_drop) if cols_to_drop else df
    
    table = doc.add_table(rows=1, cols=len(report_df.columns))
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    for i, col_name in enumerate(report_df.columns):
        hdr_cells[i].text = str(col_name)
        hdr_cells[i].paragraphs[0].runs[0].font.bold = True
        
    for index, row in report_df.iterrows():
        row_cells = table.add_row().cells
        # 🌟 FIXED: The inner data cell enumeration array population is cleanly closed
        for i, column in enumerate(report_df.columns):
            row_cells[i].text = str(row[column])
            
    filename = "Exam_Duty_Roster_Report.docx"
    doc.save(filename)
    return filename

# ----------------- LOAD EXCEL DATA (STAFF, HALLS, SUBJECTS) -----------------
try:
    staff_df = pd.read_excel("staff_list.xlsx")
    halls_df = pd.read_excel("halls_list.xlsx")
    subjects_df = pd.read_excel("subjects_list.xlsx")
except Exception as e:
    st.error(f"Excel Static Data Load Error: {e}")
    st.stop()

all_staff_names = staff_df['Name'].dropna().unique().tolist()

hall_capacities = {}
for _, row in halls_df.iterrows():
    if pd.notna(row.iloc[0]):
        h_name = str(row.iloc[0]).strip()
        h_cap = int(row.iloc[1]) if len(row) > 1 and pd.notna(row.iloc[1]) else 60
        hall_capacities[h_name] = h_cap

# ----------------- LIVE CALCULATIONS (LOAD SUMMARY) -----------------
duty_counts = {name: 0 for name in all_staff_names}
if not schedule_df.empty:
    for _, row in schedule_df.iterrows():
        sup = row['Supervisor']
        if sup in duty_counts:
            duty_counts[sup] += 1
        invs_str = str(row['Invigilators'])
        if invs_str != "nan" and invs_str != "":
            for inv in [n.strip() for n in invs_str.split(",")]:
                if inv in duty_counts:
                    duty_counts[inv] += 1

summary_data = pd.DataFrame(list(duty_counts.items()), columns=['Staff Member', 'Duties']).sort_values(by='Duties', ascending=False)

if "edit_mode" not in st.session_state:
    st.session_state.edit_mode = False
    st.session_state.edit_id = None
    st.session_state.edit_data = {}

# ----------------- LAYOUT SPLIT (SIDEBAR & MAIN) -----------------
col_control, col_display, col_summary = st.columns([1, 2.2, 0.8])

with col_control:
    if st.session_state.edit_mode:
        st.markdown("### 📝 Edit Mode Active")
        st.warning(f"Editing Record ID: {st.session_state.edit_id}")
    else:
        st.markdown("### 🛠️ Control Panel")
    
    default_date = datetime.now()
    if st.session_state.edit_mode:
        try:
            default_date = datetime.strptime(st.session_state.edit_data['Date'], "%Y-%m-%d")
        except:
            pass
    exam_date = st.date_input("Exam Date", default_date)
    date_str = str(exam_date)
    
    year_str = str(exam_date.year)
    month_str = exam_date.strftime("%B")
    
    sem_default_idx = 0
    if st.session_state.edit_mode and 'semester_key' in st.session_state.edit_data:
        for idx, opt in enumerate(["Sem 1", "Sem 2", "Sem 3"]):
            if opt.replace(" ", "") in st.session_state.edit_data['semester_key']:
                sem_default_idx = idx
    semester_opt = st.selectbox("Select Semester", ["Sem 1", "Sem 2", "Sem 3"], index=sem_default_idx)
    
    generated_semester_key = f"{year_str}-{month_str}-{semester_opt.replace(' ', '')}"
    st.caption(f"Database Partition Key: `{generated_semester_key}`")
    
    start_time = st.time_input("Start Time", time(9, 0))
    end_time = st.time_input("End Time", time(12, 0))
    time_slot = f"{start_time.strftime('%I:%M %p')} - {end_time.strftime('%I:%M %p')}"
    st.info(f"Slot: {time_slot}")
    
    sub_list = subjects_df['Subject'].tolist()
    sub_default_idx = 0
    if st.session_state.edit_mode and st.session_state.edit_data['Subject'] in sub_list:
        sub_default_idx = sub_list.index(st.session_state.edit_data['Subject'])
    selected_sub = st.selectbox("Select Subject", sub_list, index=sub_default_idx)
    
    default_students = 0
    if st.session_state.edit_mode and 'student_count' in st.session_state.edit_data:
        default_students = int(st.session_state.edit_data['student_count'])
    expected_students = st.number_input("Expected Student Count", min_value=0, value=default_students, step=1)
    
    st.write("---")
    st.markdown("#### 📍 Assign Duty")
    
    hall_list = halls_df['Hall_Name'].tolist()
    hall_default_idx = 0
    if st.session_state.edit_mode and st.session_state.edit_data['Hall'] in hall_list:
        hall_default_idx = hall_list.index(st.session_state.edit_data['Hall'])
    selected_hall = st.selectbox("Select Target Hall", hall_list, index=hall_default_idx)
    
    busy_staff_other_halls = []
    if not schedule_df.empty:
        if st.session_state.edit_mode:
            current_clashes = schedule_df[(schedule_df['Date'] == date_str) & (schedule_df['Time'] == time_slot) & (schedule_df['id'] != st.session_state.edit_id)]
        else:
            current_clashes = schedule_df[(schedule_df['Date'] == date_str) & (schedule_df['Time'] == time_slot)]
            
        busy_staff_other_halls += current_clashes['Supervisor'].tolist()
        for inv_list_str in current_clashes['Invigilators'].tolist():
            if pd.notna(inv_list_str) and inv_list_str != "" and inv_list_str != "nan":
                busy_staff_other_halls += [n.strip() for n in inv_list_str.split(",")]

    available_pool = [name for name in all_staff_names if name not in busy_staff_other_halls]

    state_key_sup = f"input_sup_{date_str}_{time_slot.replace(' ', '')}_{selected_hall}"
    state_key_inv = f"input_inv_{date_str}_{time_slot.replace(' ', '')}_{selected_hall}"
    
    if st.session_state.edit_mode:
        target_sup_val = st.session_state.edit_data.get('Supervisor', '-- Select --')
        target_inv_val = [x.strip() for x in str(st.session_state.edit_data.get('Invigilators', '')).split(",") if x.strip()]
    else:
        target_sup_val = st.session_state.get(state_key_sup, '-- Select --')
        target_inv_val = st.session_state.get(state_key_inv, [])

    filtered_supervisors = [name for name in available_pool if name not in target_inv_val]
    sup_opts = ["-- Select --"] + filtered_supervisors
    if st.session_state.edit_mode and target_sup_val not in sup_opts:
        sup_opts.append(target_sup_val)
    
    sup_default_idx = sup_opts.index(target_sup_val) if target_sup_val in sup_opts else 0
    selected_sup = st.selectbox("Supervisor", sup_opts, index=sup_default_idx, key=state_key_sup)
    
    filtered_invigilators = [name for name in available_pool if name != selected_sup]
    inv_opts = filtered_invigilators
    if st.session_state.edit_mode:
        inv_opts = list(set(inv_opts + target_inv_val))
        
    selected_invs = st.multiselect("Invigilators", inv_opts, default=[x for x in target_inv_val if x in inv_opts], key=state_key_inv)
    
    if st.session_state.edit_mode:
        col_up1, col_up2 = st.columns(2)
        with col_up1:
            if st.button("🔄 Update Assignment", use_container_width=True, type="primary"):
                if selected_sup != "-- Select --" and len(selected_invs) > 0:
                    invs_string = ", ".join(selected_invs)
                    try:
                        conn = get_db_connection()
                        cursor = conn.cursor()
                        query = """
                            UPDATE exam_assignments 
                            SET semester_key=%s, date=%s, time=%s, subject=%s, hall=%s, supervisor=%s, invigilators=%s, student_count=%s
                            WHERE id=%s
                        """
                        cursor.execute(query, (generated_semester_key, date_str, time_slot, selected_sub, selected_hall, selected_sup, invs_string, expected_students, st.session_state.edit_id))
                        cursor.close()
                        conn.close()
                        
                        st.success("Updated successfully!")
                        st.session_state.edit_mode = False
                        st.session_state.edit_id = None
                        st.rerun()
                    except Exception as ex:
                        st.error(f"DB Update Error: {ex}")
                else:
                    st.error("Select both Controller and Invigilators.")
        with col_up2:
            if st.button("❌ Cancel Edit", use_container_width=True):
                st.session_state.edit_mode = False
                st.session_state.edit_id = None
                st.rerun()
    else:
        if st.button("💾 Save Assignment", use_container_width=True, type="primary"):
            if selected_sup != "-- Select --" and len(selected_invs) > 0:
                if not schedule_df.empty and not schedule_df[(schedule_df['Date'] == date_str) & (schedule_df['Time'] == time_slot) & (schedule_df['Hall'] == selected_hall)].empty:
                    st.error(f"{selected_hall} is already allocated for this slot!")
                else:
                    invs_string = ", ".join(selected_invs)
                    try:
                        conn = get_db_connection()
                        cursor = conn.cursor()
                        query = """
                            INSERT INTO exam_assignments (semester_key, date, time, subject, hall, supervisor, invigilators, student_count)
                            VALUES (%s, %s, %s, %s, %s, %s, %s, %s)
                        """
                        cursor.execute(query, (generated_semester_key, date_str, time_slot, selected_sub, selected_hall, selected_sup, invs_string, expected_students))
                        cursor.close()
                        conn.close()
                        
                        st.success("Saved securely to Database!")
                        st.rerun()
                    except Exception as ex:
                        st.error(f"DB Insert Error: {ex}")
            else:
                st.error("Select both Controller and Invigilators.")
                
    st.write("---")
    st.markdown("#### 📥 Current Roster Report")
    
    if not schedule_df.empty:
        file_path = generate_word_report(schedule_df, "EXAMINATION DUTY ROSTER")
        with open(file_path, "rb") as file:
            st.download_button(
                label="📥 Download Word Report",
                data=file,
                file_name=f"Faculty_Exam_Roster_{generated_semester_key}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True
            )
    else:
        st.button("📥 Download Word Report", disabled=True, use_container_width=True)

# COLUMN 2: VISUAL EXAM HALL BOXES SECTIONS LAYOUT (Middle)
with col_display:
    st.markdown(f"### 🏢 Available Exam Halls View ({date_str} | {time_slot})")
    
    active_slots = schedule_df[(schedule_df['Date'] == date_str) & (schedule_df['Time'] == time_slot)] if not schedule_df.empty else pd.DataFrame()
    all_halls = halls_df['Hall_Name'].tolist()
    grid_cols = st.columns(2)
    
    for idx, hall_name in enumerate(all_halls):
        target_col = grid_cols[idx % 2]
        clean_hall_key = str(hall_name).strip()
        max_cap = hall_capacities.get(clean_hall_key, 60)
        
        with target_col:
            hall_record = active_slots[active_slots['Hall'] == clean_hall_key] if not active_slots.empty else pd.DataFrame()
            if not hall_record.empty:
                sup_assigned = hall_record.iloc[0]['Supervisor']
                invs_assigned = hall_record.iloc[0]['Invigilators']
                sub_assigned = hall_record.iloc[0]['Subject']
                date_assigned = hall_record.iloc[0]['Date']
                time_assigned = hall_record.iloc[0]['Time']
                current_students = int(hall_record.iloc[0]['student_count'] if pd.notna(hall_record.iloc[0]['student_count']) else 0)
                
                is_overcrowded = (max_cap != "N/A" and current_students > max_cap)
                box_border = "5px solid #EF4444" if is_overcrowded else "5px solid #10B981"
                box_bg = "#FEF2F2" if is_overcrowded else "#F0FDF4"
                badge_class = "badge-danger" if is_overcrowded else "badge-success"
                title_color = "#DC2626" if is_overcrowded else "#1E40AF"
                title_prefix = "⚠️" if is_overcrowded else "🟢"
                
                st.markdown(f"""
                    <div class="hall-box" style="border-left: {box_border}; background-color: {box_bg};">
                        <span class="{badge_class}">Students: {current_students} / {max_cap}</span>
                        <div class="hall-title" style="color: {title_color};">{title_prefix} {clean_hall_key}</div>
                        <div class="hall-details">
                            <b>Schedule:</b> {date_assigned} ({time_assigned})<br>
                            <b>Subject:</b> {sub_assigned}<br>
                            <b>Supervisor:</b> {sup_assigned}<br>
                            <b>Invigilators:</b> {invs_assigned}
                        </div>
                    </div>
                """, unsafe_allow_html=True)
            else:
                st.markdown(f"""
                    <div class="hall-box" style="border-left: 5px solid #64748B; background-color: #F8FAFC;">
                        <span class="badge-capacity">Max: {max_cap}</span>
                        <div class="hall-title" style="color:#64748B;">⚪ {clean_hall_key}</div>
                        <div class="hall-details" style="color:#94A3B8;">
                            <i>No exam scheduled or staff assigned for this time slot.</i>
                        </div>
                    </div>
                """, unsafe_allow_html=True)

# COLUMN 3: LOAD SUMMARY LIST (Right Side)
with col_summary:
    st.markdown("### 📊 Summary")
    st.markdown("<div class='summary-title'>Load Summary</div>", unsafe_allow_html=True)
    st.dataframe(summary_data, use_container_width=True, hide_index=True, height=520)

st.write("---")

# ----------------- LIVE DATA EDITOR PANEL -----------------
st.markdown("### 📝 Active Assignments Management Panel (Live Editor)")
if not schedule_df.empty:
    st.caption("Here are all the live entries stored inside the database. Use the action buttons below to fix entry errors.")
    
    for index, row in schedule_df.iterrows():
        with st.container():
            r_col1, r_col2, r_col3, r_col4, r_col5 = st.columns([1.2, 1.2, 1.5, 2.5, 1.2])
            with r_col1:
                st.write(f"📅 **{row['Date']}**")
                st.caption(f"ID: {row['id']} | {row['semester_key']}")
            with r_col2:
                st.write(f"⏰ {row['Time']}")
            with r_col3:
                st.write(f"🏢 {row['Hall']} (Allocated: {row['student_count']})")
                st.caption(f"📚 {row['Subject']}")
            with r_col4:
                st.write(f"👤 **Sup:** {row['Supervisor']}")
                st.write(f"👥 **Invs:** {row['Invigilators']}")
            with r_col5:
                btn_col1, btn_col2 = st.columns(2)
                with btn_col1:
                    if st.button("✏️ Edit", key=f"edit_{row['id']}", use_container_width=True):
                        st.session_state.edit_mode = True
                        st.session_state.edit_id = int(row['id'])
                        st.session_state.edit_data = row.to_dict()
                        st.rerun()
                with btn_col2:
                    if st.button("🗑️", key=f"del_{row['id']}", use_container_width=True, help="Delete assignment"):
                        try:
                            conn = get_db_connection()
                            cursor = conn.cursor()
                            cursor.execute("DELETE FROM exam_assignments WHERE id=%s", (int(row['id']),))
                            cursor.close()
                            conn.close()
                            st.toast(f"Record {row['id']} Deleted!", icon="🗑️")
                            st.rerun()
                        except Exception as ex:
                            st.error(f"Error deleting: {ex}")
            st.markdown("<hr style='margin:0.5em 0px; border-color:#E2E8F0;'/>", unsafe_allow_html=True)
else:
    st.info("No active logs stored in the database yet.")
